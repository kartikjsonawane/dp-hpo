"""
fill_results.py — reads results_v2.json, fills the 4 results tables
in DP_HPO_V2_Paper.docx, and saves DP_HPO_V2_Paper_FINAL.docx.

Usage:
    python fill_results.py
    python fill_results.py --results results_v2.json --input DP_HPO_V2_Paper.docx

Requires:  pip install python-docx scipy numpy
"""

import argparse
import json
import sys
import os
import numpy as np
from scipy.stats import wilcoxon
from copy import deepcopy

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    sys.exit("ERROR: pip install python-docx")

# ── Constants ─────────────────────────────────────────────────────────────────

METHODS_ORDER = [
    "Grid Search",
    "Random Search (k=20)",
    "Random Search (k=10)",
    "Bayesian Opt",
    "Optuna (TPE)",
    "Hyperband",
    "BOHB",
    "SMAC",
    "DP-HPO (Proposed)",
]
EVALS = {
    "Grid Search": "108",
    "Random Search (k=20)": "20",
    "Random Search (k=10)": "10",
    "Bayesian Opt": "20",
    "Optuna (TPE)": "20",
    "Hyperband": "30",
    "BOHB": "30",
    "SMAC": "20",
    "DP-HPO (Proposed)": "10",
}

# These are the dataset keys in main_results — must match run_experiments.py
DATASET_KEYS = [
    "UCI Breast Cancer",
    "MNIST",
    "Fashion-MNIST",
    "Adult Income",
]

# The placeholder text used in the docx (one table per dataset = 4 tables)
PLACEHOLDER = "[pending]"

ALPHA = 0.05

# ── Load results ──────────────────────────────────────────────────────────────

def load_results(path):
    with open(path) as f:
        data = json.load(f)
    return data["main_results"]

# ── Wilcoxon + Bonferroni ─────────────────────────────────────────────────────

def compute_stats(main_results, bonferroni_m):
    """
    Returns dict: dataset → method → {mean_pct, std_pct, pval, sig}
    """
    alpha_adj = ALPHA / bonferroni_m
    out = {}
    for ds_key, ds_data in main_results.items():
        dp_accs = np.array(ds_data["DP-HPO (Proposed)"]["accs"])
        out[ds_key] = {}
        for method in METHODS_ORDER:
            if method not in ds_data:
                print(f"  WARNING: method '{method}' missing in {ds_key}")
                out[ds_key][method] = {"mean_pct": None, "std_pct": None, "pval": None, "sig": False}
                continue
            accs = np.array(ds_data[method]["accs"])
            mean_pct = float(np.mean(accs) * 100)
            std_pct  = float(np.std(accs)  * 100)
            if method == "DP-HPO (Proposed)":
                pval = None
                sig  = False
            else:
                try:
                    _, pval = wilcoxon(dp_accs, accs, alternative="two-sided")
                    sig = pval < alpha_adj
                except Exception:
                    pval = None
                    sig  = False
            out[ds_key][method] = {
                "mean_pct": mean_pct,
                "std_pct":  std_pct,
                "pval":     pval,
                "sig":      sig,
            }
    return out, alpha_adj

# ── Identify best non-DP-HPO method ──────────────────────────────────────────

def best_baseline(ds_stats):
    best_method = None
    best_mean   = -1
    for method in METHODS_ORDER[:-1]:          # exclude DP-HPO
        m = ds_stats[method]["mean_pct"]
        if m is not None and m > best_mean:
            best_mean   = m
            best_method = method
    return best_method

# ── Find "pending" tables in docx ────────────────────────────────────────────

def find_results_tables(doc):
    """
    Return the 4 results tables (those containing '[pending]' cells).
    """
    tables = []
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                if PLACEHOLDER in cell.text:
                    tables.append(tbl)
                    break
            else:
                continue
            break
    return tables

# ── Fill one table ────────────────────────────────────────────────────────────

def fill_table(table, ds_stats, alpha_adj):
    """
    Replace [pending] cells in a results table.
    Row order assumed: header + one row per method (9 methods).
    Columns: Method | Mean Acc (%) | Std (%) | Evals | p-value
    """
    best = best_baseline(ds_stats)
    dp_mean = ds_stats["DP-HPO (Proposed)"]["mean_pct"]

    method_rows = [r for r in table.rows if PLACEHOLDER in r.cells[0].text
                   or any(PLACEHOLDER in c.text for c in r.cells)]

    for row in method_rows:
        # Identify method from first cell
        method_cell_text = row.cells[0].text.strip()
        matched = None
        for m in METHODS_ORDER:
            if m in method_cell_text:
                matched = m
                break
        if matched is None:
            continue

        s = ds_stats.get(matched)
        if s is None:
            continue

        mean_pct = s["mean_pct"]
        std_pct  = s["std_pct"]
        pval     = s["pval"]
        sig      = s["sig"]

        # Col 1: Mean Acc
        cell_mean = row.cells[1]
        _set_cell(cell_mean, f"{mean_pct:.2f}",
                  bold=(matched == best or matched == "DP-HPO (Proposed)"))

        # Col 2: Std
        cell_std = row.cells[2]
        _set_cell(cell_std, f"{std_pct:.2f}")

        # Col 3: Evals (already filled but may have placeholder)
        cell_evals = row.cells[3]
        if PLACEHOLDER in cell_evals.text:
            _set_cell(cell_evals, EVALS.get(matched, "—"))

        # Col 4: p-value
        cell_p = row.cells[4]
        if matched == "DP-HPO (Proposed)":
            _set_cell(cell_p, "—")
        elif pval is None:
            _set_cell(cell_p, "N/A")
        else:
            pval_str = f"{pval:.4f}" if pval >= 0.0001 else "<0.0001"
            _set_cell(cell_p, pval_str, bold=sig,
                      color=(0x00, 0x70, 0x00) if sig else None)

# ── Cell helpers ──────────────────────────────────────────────────────────────

def _set_cell(cell, text, bold=False, color=None):
    """Clear cell and write text with optional bold/color."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    # Write to first paragraph
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = "Times New Roman"
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)

# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Fill DP-HPO manuscript tables from results_v2.json")
    parser.add_argument("--results", default="results_v2.json",
                        help="Path to results_v2.json (default: results_v2.json)")
    parser.add_argument("--input",   default="DP_HPO_V2_Paper.docx",
                        help="Source .docx (default: DP_HPO_V2_Paper.docx)")
    parser.add_argument("--output",  default="DP_HPO_V2_Paper_FINAL.docx",
                        help="Output .docx (default: DP_HPO_V2_Paper_FINAL.docx)")
    args = parser.parse_args()

    # Resolve paths relative to this script's directory
    script_dir = os.path.dirname(os.path.abspath(__file__))
    results_path = os.path.join(script_dir, args.results)
    input_path   = os.path.join(script_dir, args.input)
    output_path  = os.path.join(script_dir, args.output)

    print(f"Loading results: {results_path}")
    if not os.path.exists(results_path):
        sys.exit(f"ERROR: {results_path} not found. Run python run_experiments.py first.")

    main_results = load_results(results_path)

    # Remap dataset keys — run_experiments.py may use short or long keys
    # Try to map by substring
    available_keys = list(main_results.keys())
    print(f"Datasets in results_v2.json: {available_keys}")

    bonferroni_m = 8 * len(available_keys)   # 8 baselines × datasets
    print(f"Bonferroni m = {bonferroni_m}   α_adj = {ALPHA/bonferroni_m:.5f}")

    stats, alpha_adj = compute_stats(main_results, bonferroni_m)

    # Print summary
    print("\n── Summary ────────────────────────────────────────────────────")
    for ds_key, ds_stats in stats.items():
        print(f"\n  {ds_key}")
        print(f"  {'Method':<30} {'Mean%':>8} {'Std%':>6} {'p-value':>10} {'Sig':>5}")
        print("  " + "-"*60)
        for method in METHODS_ORDER:
            s = ds_stats.get(method, {})
            m = s.get("mean_pct")
            sd = s.get("std_pct")
            p  = s.get("pval")
            sig = s.get("sig", False)
            m_str  = f"{m:.2f}" if m is not None else "—"
            sd_str = f"{sd:.2f}" if sd is not None else "—"
            p_str  = f"{p:.4f}" if p is not None else "—"
            sig_str = "✓" if sig else ""
            print(f"  {method:<30} {m_str:>8} {sd_str:>6} {p_str:>10} {sig_str:>5}")

    # Open docx
    print(f"\nOpening: {input_path}")
    doc = Document(input_path)

    results_tables = find_results_tables(doc)
    print(f"Found {len(results_tables)} results table(s) with '[pending]' cells")

    if len(results_tables) == 0:
        print("WARNING: No pending tables found. Is the manuscript already filled?")
    else:
        # Match tables to datasets by order (Table IV → V → VI → VII)
        ds_keys = list(stats.keys())
        for i, tbl in enumerate(results_tables):
            if i >= len(ds_keys):
                print(f"  Table {i+1}: no matching dataset (only {len(ds_keys)} datasets)")
                break
            ds_key = ds_keys[i]
            print(f"  Filling Table {i+1} → {ds_key}")
            fill_table(tbl, stats[ds_key], alpha_adj)

    doc.save(output_path)
    print(f"\nSaved: {output_path}")

    # Count significant results
    total_sig = sum(
        s["sig"]
        for ds_stats in stats.values()
        for method, s in ds_stats.items()
        if method != "DP-HPO (Proposed)" and s.get("pval") is not None
    )
    total_tests = sum(
        1
        for ds_stats in stats.values()
        for method, s in ds_stats.items()
        if method != "DP-HPO (Proposed)" and s.get("pval") is not None
    )
    print(f"\nSignificant results (Bonferroni α_adj={alpha_adj:.5f}): "
          f"{total_sig}/{total_tests} comparisons")
    print("Done.")

if __name__ == "__main__":
    main()
